"""
PDF ingestion pipeline for Intelli-Credit.

Extraction strategy:
1) PyMuPDF (fitz) for fast native text
2) pdfplumber for tables/line structures
3) Qwen2.5-VL OCR for scanned/image-only pages (with local fallback)
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Sequence

import fitz
import pdfplumber
from pydantic import ValidationError

from backend.core.ingestion.qwen_vl_ocr import QwenVLOCR
from backend.schemas.credit import (
    BankLimit,
    DocumentType,
    ExtractedFinancialData,
    LegalDispute,
    MoneyAmount,
)


def parse_revenue_to_crore(raw_value: str | float | int, unit_hint: str = "") -> float:
    """
    Normalize monetary values to crore units.
    """
    clean = str(raw_value).replace("\u20b9", "").replace(",", "").strip()
    try:
        value = float(clean)
    except (TypeError, ValueError):
        return 0.0

    unit_low = (unit_hint or "").lower()
    if "lakh" in unit_low or "lac" in unit_low:
        return value / 100.0
    if "million" in unit_low or "mn" in unit_low:
        return value / 10.0
    if "crore" in unit_low or " cr" in f" {unit_low}" or unit_low.strip() == "cr":
        return value

    # Magnitude fallback
    if value > 100000:
        return value / 100.0
    if value > 10000:
        return value / 10.0
    return value


def extract_revenue(text: str) -> Optional[float]:
    """
    Extract annual revenue in crore from common Indian financial document formats.
    """
    match = re.search(
        r"[Rr]evenue[\s\S]{0,160}?₹\s*([\d,]+(?:\.\d+)?)\s*(?:[Cc]r(?:ores?)?)\b",
        text,
    )
    if match:
        return float(match.group(1).replace(",", ""))

    match = re.search(
        r"[Rr]evenue[^0-9]{0,80}([\d,]+(?:\.\d+)?)\s*[Ll]akhs?\b",
        text,
    )
    if match:
        lakhs = float(match.group(1).replace(",", ""))
        return lakhs / 100.0

    match = re.search(
        r"[Gg]ross\s+[Rr]eceipts[^0-9]*([\d,]+(?:\.\d+)?)",
        text,
    )
    if match:
        return float(match.group(1).replace(",", ""))

    return None


DOC_KEYWORDS = {
    DocumentType.ANNUAL_REPORT: [
        "directors' report",
        "board of directors",
        "standalone financial",
    ],
    DocumentType.SANCTION_LETTER: [
        "sanctioned amount",
        "rate of interest",
        "security",
    ],
    DocumentType.LEGAL_NOTICE: ["cause of action", "legal notice", "advocate"],
    DocumentType.RATING_REPORT: ["crisil", "icra", "care ratings", "credit rating"],
    DocumentType.FINANCIAL_STATEMENT: [
        "balance sheet",
        "profit & loss",
        "cash flow",
    ],
    DocumentType.BOARD_MINUTES: [
        "minutes of meeting",
        "resolution passed",
        "quorum",
    ],
}


@dataclass
class ParseResult:
    text: str
    method: str
    confidence: float


class IntelliCreditPDFParser:
    """
    Parse a PDF into structured, credit-oriented financial entities.
    """

    min_text_threshold: int = 120

    def __init__(self) -> None:
        self.ocr = QwenVLOCR()

    def parse(self, file_path: str) -> ExtractedFinancialData:
        suffix = Path(file_path).suffix.lower()
        if suffix in {".jpg", ".jpeg", ".png"}:
            parse_result = self._extract_text_from_image(file_path)
        elif suffix == ".docx":
            parse_result = self._extract_text_from_docx(file_path)
        else:
            parse_result = self._extract_text(file_path)
        doc_type = self._detect_document_type(parse_result.text)

        extraction = self._extract_entities(parse_result.text, doc_type)
        extraction.document_type = doc_type
        extraction.extraction_confidence = min(
            1.0, max(0.1, (extraction.extraction_confidence + parse_result.confidence) / 2)
        )

        if extraction.extraction_confidence < 0.7:
            extraction.needs_human_review = True
            extraction.key_risks_mentioned.append(
                "Low extraction confidence (<0.7). Manual credit analyst review required."
            )

        if extraction.auditor_qualifications:
            extraction.key_risks_mentioned.append(
                "CRITICAL: Auditor qualifications present in submitted document."
            )
        if extraction.going_concern_doubts:
            extraction.key_risks_mentioned.append(
                "CRITICAL: Going concern doubt observed in auditor commentary."
            )

        return extraction

    def _extract_text(self, file_path: str) -> ParseResult:
        """
        Multi-stage extraction with confidence approximation.
        """
        fitz_text = self._extract_with_fitz(file_path)
        if len(fitz_text) >= self.min_text_threshold:
            return ParseResult(text=fitz_text, method="fitz", confidence=0.9)

        plumber_text = self._extract_with_pdfplumber(file_path)
        combined = "\n".join([fitz_text, plumber_text]).strip()
        if len(combined) >= self.min_text_threshold:
            return ParseResult(text=combined, method="fitz+pdfplumber", confidence=0.78)

        ocr_result = self.ocr.extract_text_from_pdf(file_path)
        final_text = "\n".join([combined, ocr_result.text]).strip()
        confidence = ocr_result.confidence if final_text else 0.3
        if len(final_text) < self.min_text_threshold:
            confidence = min(confidence, 0.45)
        return ParseResult(text=final_text, method=ocr_result.method, confidence=confidence)

    def _extract_text_from_image(self, file_path: str) -> ParseResult:
        ocr_result = self.ocr.extract_text_from_image_path(file_path)
        confidence = ocr_result.confidence
        if len(ocr_result.text) < self.min_text_threshold:
            confidence = min(confidence, 0.45)
        return ParseResult(text=ocr_result.text, method=ocr_result.method, confidence=confidence)

    def _extract_text_from_docx(self, file_path: str) -> ParseResult:
        """
        Extract text from a DOCX file (paragraphs + table cells).
        """
        try:
            from docx import Document as WordDocument

            doc = WordDocument(file_path)
            chunks: List[str] = []

            for paragraph in doc.paragraphs:
                text = (paragraph.text or "").strip()
                if text:
                    chunks.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                        row_cells.append(cell_text)
                    if any(row_cells):
                        chunks.append(" | ".join(row_cells))

            joined = "\n".join(chunks).strip()
            confidence = 0.82 if len(joined) >= self.min_text_threshold else 0.58
            return ParseResult(text=joined, method="python-docx", confidence=confidence)
        except Exception:
            return ParseResult(text="", method="python-docx", confidence=0.2)

    def _extract_with_fitz(self, file_path: str) -> str:
        text_chunks: List[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_chunks.append(page.get_text("text") or "")  # type: ignore[reportArgumentType]
        return "\n".join(text_chunks).strip()

    def _extract_with_pdfplumber(self, file_path: str) -> str:
        text_chunks: List[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
                table = page.extract_table()
                if table:
                    flattened = [" | ".join([str(cell or "") for cell in row]) for row in table]
                    text_chunks.append("\n".join(flattened))
        return "\n".join(text_chunks).strip()

    def _detect_document_type(self, text: str) -> DocumentType:
        low = text.lower()
        scored = []
        for doc_type, keywords in DOC_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in low)
            scored.append((doc_type, score))
        best, best_score = max(scored, key=lambda x: x[1])
        return best if best_score > 0 else DocumentType.UNKNOWN

    def _extract_entities(self, text: str, doc_type: DocumentType) -> ExtractedFinancialData:
        cin = self._find_first(r"\b([LU]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6})\b", text)
        pan = self._find_first(r"\b([A-Z]{5}[0-9]{4}[A-Z])\b", text)
        company_name = self._extract_company_name(text)

        revenue = []
        extracted_revenue_cr = extract_revenue(text)
        if extracted_revenue_cr is not None and extracted_revenue_cr > 0:
            revenue = [
                MoneyAmount(
                    amount=round(extracted_revenue_cr, 2),
                    period="annual_1",
                    source="pdf",
                    confidence_score=0.82,
                )
            ]
        else:
            revenue = self._extract_money_series(
                text,
                patterns=[
                    r"revenue(?: from operations)?[\s\S]{0,120}?₹\s*([\d,]+(?:\.\d+)?)\s*(crores?|cr)?",
                    r"revenue(?: from operations)?[\s\S]{0,120}?([\d,]+(?:\.\d+)?)\s*(crores?|cr|lakhs?|lacs?|million|mn)?",
                    r"annual turnover[^\d]{0,40}([\d,]+(?:\.\d+)?)\s*(crores?|cr|lakhs?|lacs?|million|mn)?",
                    r"gross receipts[^\d]{0,40}([\d,]+(?:\.\d+)?)\s*(crores?|cr|lakhs?|lacs?|million|mn)?",
                ],
                period_hint="annual",
                source="pdf",
            )
        profits = self._extract_money_series(
            text,
            patterns=[
                r"(?:pat|profit after tax|net profit)[^\d]{0,30}([\d,]+(?:\.\d+)?)\s*(crores?|cr|lakhs?|lacs?|million|mn)?"
            ],
            period_hint="annual",
            source="pdf",
        )
        total_debt_match = self._find_first(
            r"(?:total debt|borrowings)[^\d]{0,20}([\d,]+(?:\.\d+)?)", text
        )
        total_debt = (
            MoneyAmount(
                amount=float(total_debt_match.replace(",", "")),
                period="latest",
                source="pdf",
                confidence_score=0.75,
            )
            if total_debt_match
            else None
        )
        customer_concentration_match = self._find_first(
            r"(?:top\s*5\s*customers?|customer concentration(?:\s*risk)?)[^\d%]{0,40}(\d{1,3}(?:\.\d+)?)\s*%",
            text,
        )
        customer_concentration_top5_pct = (
            float(customer_concentration_match) if customer_concentration_match else None
        )
        crisil_rating_action_match = self._find_first(
            r"\b(upgraded|downgraded|reaffirmed|assigned)\b",
            text,
        )
        crisil_rating_action = (
            str(crisil_rating_action_match).upper() if crisil_rating_action_match else None
        )
        crisil_rating_match = self._find_first(
            r"(?:crisil|icra|care(?:\s+ratings)?)[\s\S]{0,60}?\b([A-Z]{1,4}[+-]?)\b",
            text,
        )
        crisil_rating = str(crisil_rating_match).upper() if crisil_rating_match else None

        strategic_partner_match = self._find_first(
            r"(?:strategic partner(?:ship)?|collaboration|alliance|mou)[\s\S]{0,80}?\b([A-Z][A-Za-z0-9&.,()\- ]{2,80})\b",
            text,
        )
        strategic_partnership_mentioned = (
            " ".join(str(strategic_partner_match).split()) if strategic_partner_match else None
        )

        legal_disputes = self._extract_legal_disputes(text)
        collateral = self._extract_collateral(text)
        bank_limits = self._extract_bank_limits(text)
        related_party = self._extract_related_party_transactions(text)
        auditor_quals = self._extract_auditor_qualifications(text)
        going_concern = bool(re.search(r"going concern", text, flags=re.IGNORECASE))

        confidence = 0.55
        if cin:
            confidence += 0.05
        if pan:
            confidence += 0.05
        if revenue:
            confidence += 0.1
        if profits:
            confidence += 0.1
        if doc_type != DocumentType.UNKNOWN:
            confidence += 0.1

        try:
            return ExtractedFinancialData(
                company_name=company_name or "Unknown Company",
                cin_number=cin,
                pan_number=pan,
                crisil_rating=crisil_rating,
                crisil_rating_action=crisil_rating_action,
                customer_concentration_top5_pct=customer_concentration_top5_pct,
                strategic_partnership_mentioned=strategic_partnership_mentioned,
                revenue_figures=revenue,
                profit_figures=profits,
                total_debt=total_debt,
                existing_bank_limits=bank_limits,
                collateral_descriptions=collateral,
                legal_disputes=legal_disputes,
                key_risks_mentioned=[],
                related_party_transactions=related_party,
                auditor_qualifications=auditor_quals,
                going_concern_doubts=going_concern,
                document_type=doc_type,
                extraction_confidence=round(min(confidence, 0.95), 3),
                needs_human_review=False,
            )
        except ValidationError:
            return ExtractedFinancialData(
                company_name=company_name or "Unknown Company",
                document_type=doc_type,
                extraction_confidence=0.4,
                needs_human_review=True,
            )

    @staticmethod
    def _find_first(pattern: str, text: str) -> Optional[str]:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if not match:
            return None
        return match.group(1) if match.groups() else match.group(0)

    @staticmethod
    def _extract_company_name(text: str) -> Optional[str]:
        patterns = [
            r"(?i)([A-Z][A-Za-z0-9&.,\-\s]{3,80}(?:Pvt\.?\s*Ltd\.?|Limited|LLP))",
            r"(?i)name of the company[:\s]+([A-Za-z0-9&.,\-\s]{4,100})",
        ]
        for pat in patterns:
            m = re.search(pat, text)
            if m:
                return " ".join(m.group(1).split())
        return None

    @staticmethod
    def _extract_money_series(
        text: str,
        *,
        patterns: Sequence[str],
        period_hint: str,
        source: str,
    ) -> List[MoneyAmount]:
        amounts: List[MoneyAmount] = []
        for pattern in patterns:
            matches = re.findall(pattern, text, flags=re.IGNORECASE)
            for idx, m in enumerate(matches[:5], start=1):
                raw: str
                unit_hint = ""
                if isinstance(m, tuple):
                    raw = str(m[0])
                    if len(m) > 1:
                        unit_hint = str(m[1] or "")
                else:
                    raw = str(m)
                val = parse_revenue_to_crore(raw, unit_hint)
                if val <= 0:
                    continue
                amounts.append(
                    MoneyAmount(
                        amount=round(val, 2),
                        period=f"{period_hint}_{idx}",
                        source=source,
                        confidence_score=0.72,
                    )
                )
        return amounts

    @staticmethod
    def _extract_collateral(text: str) -> List[str]:
        lines = []
        for line in text.splitlines():
            low = line.lower()
            if "collateral" in low or "security" in low or "hypothecation" in low:
                cleaned = " ".join(line.split())
                if len(cleaned) > 12:
                    lines.append(cleaned[:200])
        return lines[:8]

    @staticmethod
    def _extract_bank_limits(text: str) -> List[BankLimit]:
        limits: List[BankLimit] = []
        pattern = re.compile(
            r"([A-Za-z&\s]{3,40}bank)[^\d]{0,20}([\d,]+(?:\.\d+)?)",
            flags=re.IGNORECASE,
        )
        for lender, amount in pattern.findall(text):
            try:
                value = float(amount.replace(",", ""))
            except ValueError:
                continue
            limits.append(
                BankLimit(
                    lender=" ".join(lender.split()),
                    limit_amount=value,
                    facility_type="Working Capital",
                    confidence_score=0.7,
                )
            )
        return limits[:10]

    @staticmethod
    def _extract_legal_disputes(text: str) -> List[LegalDispute]:
        disputes: List[LegalDispute] = []
        for sentence in re.split(r"[.\n]", text):
            low = sentence.lower()
            if any(key in low for key in ("court", "petition", "litigation", "notice", "dispute")):
                cleaned = " ".join(sentence.split())
                if len(cleaned) < 20:
                    continue
                disputes.append(
                    LegalDispute(
                        title=cleaned[:180],
                        status="open",
                        confidence_score=0.66,
                    )
                )
        return disputes[:8]

    @staticmethod
    def _extract_related_party_transactions(text: str) -> List[str]:
        findings = []
        relationship_keywords = (
            "related party",
            "subsidiary",
            "associate",
            "joint venture",
            "wholly owned",
            "strategic partner",
            "strategic partnership",
            "collaboration",
            "alliance",
            "mou",
            "rating agency",
            "rated by",
            "crisil",
            "icra",
            "care ratings",
            "fitch",
            "moody",
        )
        for sentence in re.split(r"[.\n]", text):
            low = sentence.lower()
            if any(keyword in low for keyword in relationship_keywords):
                clean = " ".join(sentence.split())
                if clean and clean not in findings:
                    findings.append(clean[:240])
        return findings[:20]

    @staticmethod
    def _extract_auditor_qualifications(text: str) -> List[str]:
        findings = []
        pattern = re.compile(
            r"(qualified opinion|emphasis of matter|material uncertainty|adverse opinion|disclaimer)",
            flags=re.IGNORECASE,
        )
        for sentence in re.split(r"[.\n]", text):
            if pattern.search(sentence):
                clean = " ".join(sentence.split())
                if clean:
                    findings.append(clean[:220])
        return findings[:8]
