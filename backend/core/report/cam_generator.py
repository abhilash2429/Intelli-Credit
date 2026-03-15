"""
DOCX CAM generator.
"""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from backend.config import settings
from backend.core.formatting import format_currency_cr, format_percentage, format_ratio
from backend.core.report.five_c_analyzer import analyze_five_cs
from backend.core.report.templates.cam_template import CAM_SECTION_ORDER


class CAMGenerator:
    """
    Create a professional Credit Appraisal Memorandum as .docx.
    """

    def __init__(self, output_dir: str | None = None) -> None:
        self.output_dir = Path(output_dir or settings.cam_output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        *,
        company: Dict[str, Any],
        decision: Dict[str, Any],
        explanation: Dict[str, Any],
        research_findings: List[Dict[str, Any]],
        research_narrative: str | None = None,
        features: Dict[str, float],
        cross_validation: Dict[str, Any],
        due_diligence: Dict[str, Any] | None = None,
        borrower_context: Dict[str, Any] | None = None,
    ) -> str:
        doc = Document()
        due_diligence = due_diligence or {}
        borrower_context = borrower_context or {}
        self._apply_styles(doc)
        self._header(doc)
        self._toc_placeholder(doc)
        self._section_executive_summary(doc, company, decision)
        self._section_borrower_profile(doc, company, due_diligence, borrower_context)
        self._section_five_cs(doc, features)
        self._section_financial_analysis(doc, features, cross_validation)
        self._section_risk_assessment(doc, explanation, cross_validation)
        self._section_research(doc, research_findings, research_narrative)
        self._section_recommendation(doc, decision, explanation, due_diligence)
        self._section_mitigants(doc, cross_validation, features)
        self._section_annexures(doc, research_findings, explanation, due_diligence, decision)
        self._add_footer_page_number(doc)

        filename = f"CAM_{company.get('name', 'company').replace(' ', '_')}_{date.today().isoformat()}.docx"
        path = self.output_dir / filename
        doc.save(path)  # type: ignore[reportArgumentType]
        return str(path)

    @staticmethod
    def _apply_styles(doc: Document) -> None:  # type: ignore[reportGeneralTypeIssues]
        if "Heading 1" in doc.styles:
            h1 = doc.styles["Heading 1"]
            h1.font.color.rgb = RGBColor(15, 44, 88)
            h1.font.size = Pt(14)
            h1.font.bold = True

        if "Heading 2" in doc.styles:
            h2 = doc.styles["Heading 2"]
            h2.font.color.rgb = RGBColor(22, 67, 122)
            h2.font.size = Pt(12)
            h2.font.bold = True

    @staticmethod
    def _header(doc: Document) -> None:  # type: ignore[reportGeneralTypeIssues]
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run("CREDIT APPRAISAL MEMORANDUM")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(10, 38, 90)

        p2 = doc.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p2.add_run(f"{settings.report_bank_name} | Corporate & Institutional Banking").bold = True
        p3 = doc.add_paragraph("CONFIDENTIAL — CREDIT COMMITTEE")
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p3.runs[0].font.color.rgb = RGBColor(120, 0, 0)
        doc.add_paragraph(f"Date: {date.today().isoformat()} | Logo: [BANK LOGO PLACEHOLDER]")

    @staticmethod
    def _toc_placeholder(doc: Document) -> None:  # type: ignore[reportGeneralTypeIssues]
        doc.add_heading("Table of Contents", level=1)
        for section in CAM_SECTION_ORDER:
            doc.add_paragraph(section)
        doc.add_page_break()

    @staticmethod
    def _section_executive_summary(doc: Document, company: Dict[str, Any], decision: Dict[str, Any]) -> None:  # type: ignore[reportGeneralTypeIssues]
        doc.add_heading("SECTION 1: EXECUTIVE SUMMARY", level=1)
        doc.add_paragraph(
            f"Company: {company.get('name')} | CIN: {company.get('cin', 'N/A')} | Sector: {company.get('sector')}"
        )
        doc.add_paragraph(
            f"Loan Request: {format_currency_cr(company.get('loan_amount_requested', 0))} | "
            f"Tenor: {company.get('loan_tenor_months', 36)} months | Purpose: {company.get('loan_purpose', 'N/A')}"
        )
        box = doc.add_paragraph()
        recommendation = decision.get('recommendation', 'PENDING')
        if recommendation == "REJECT":
            box.add_run(
                f"Recommendation: REJECT | "
                f"Limit: Not Sanctioned (Requested: {format_currency_cr(decision.get('recommended_loan_amount', 0))}) | "
                f"Rate: N/A"
            ).bold = True
            if decision.get("rule_hits"):
                doc.add_paragraph(f"Rejection reasons: {', '.join(decision['rule_hits'])}")
        else:
            box.add_run(
                f"Recommendation: {recommendation} | "
                f"Amount: {format_currency_cr(decision.get('recommended_loan_amount', 0))} | "
                f"Rate: {decision.get('recommended_interest_rate', 0):.2f}% p.a."
            ).bold = True

    @staticmethod
    def _section_borrower_profile(
        doc: Document,  # type: ignore[reportGeneralTypeIssues]
        company: Dict[str, Any],
        due_diligence: Dict[str, Any],
        borrower_context: Dict[str, Any],
    ) -> None:
        doc.add_heading("SECTION 2: BORROWER PROFILE", level=1)
        doc.add_paragraph(
            f"Borrower: {company.get('name', 'N/A')} | Sector: {company.get('sector', 'N/A')} | "
            f"CIN: {company.get('cin', 'N/A')}."
        )
        business_highlights = (
            borrower_context.get("borrower_business_highlights")
            or borrower_context.get("business_highlights")
            or "No additional borrower business highlights submitted."
        )
        major_customers = (
            borrower_context.get("borrower_major_customers")
            or borrower_context.get("major_customers")
            or "Not disclosed in borrower clarification."
        )
        doc.add_paragraph(f"Business profile (borrower submission): {business_highlights}")
        doc.add_paragraph(f"Major customers / revenue anchors: {major_customers}")

        doc.add_heading("Borrower Clarifications (Finance Officer Input)", level=2)
        officer_name = borrower_context.get("borrower_finance_officer_name") or borrower_context.get(
            "finance_officer_name"
        )
        officer_role = borrower_context.get("borrower_finance_officer_role") or borrower_context.get(
            "finance_officer_role"
        )
        if officer_name:
            doc.add_paragraph(f"Primary respondent: {officer_name} ({officer_role or 'Role not specified'})")
        else:
            doc.add_paragraph("Primary respondent: Not provided by borrower.")

        highlights = (
            borrower_context.get("borrower_business_highlights")
            or borrower_context.get("business_highlights")
            or "No additional borrower business highlights submitted."
        )
        risks = (
            borrower_context.get("borrower_disclosed_risks")
            or borrower_context.get("disclosed_risks")
            or "No borrower-declared risks submitted."
        )
        doc.add_paragraph(f"Business highlights: {highlights}")
        doc.add_paragraph(f"Borrower-declared risks: {risks}")
        if due_diligence.get("factory_capacity_utilization") is not None:
            doc.add_paragraph(
                f"Declared capacity utilization: {format_percentage(due_diligence.get('factory_capacity_utilization'))}"  # type: ignore[reportArgumentType]
            )
        if due_diligence.get("management_integrity_score") is not None:
            doc.add_paragraph(
                "Management integrity score (from due diligence input): "
                f"{format_ratio(due_diligence.get('management_integrity_score'), decimals=1, suffix='/10')}"  # type: ignore[reportArgumentType]
            )

    @staticmethod
    def _section_five_cs(doc: Document, features: Dict[str, float]) -> None:  # type: ignore[reportGeneralTypeIssues]
        doc.add_heading("SECTION 3: THE FIVE Cs ANALYSIS", level=1)
        five_c = analyze_five_cs(features)
        labels = [
            ("Character", "character"),
            ("Capacity", "capacity"),
            ("Capital", "capital"),
            ("Collateral", "collateral"),
            ("Conditions", "conditions"),
        ]
        for title, key in labels:
            item = five_c[key]
            doc.add_heading(f"3.{labels.index((title, key))+1} {title}", level=2)
            doc.add_paragraph(f"Score: {item['score']}/10 | Risk Level: {item['risk_level']}")

    @staticmethod
    def _section_financial_analysis(
        doc: Document,  # type: ignore[reportGeneralTypeIssues]
        features: Dict[str, float],
        cross_validation: Dict[str, Any],
    ) -> None:
        doc.add_heading("SECTION 4: FINANCIAL ANALYSIS", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"

        for metric, value in [
            ("DSCR", features.get("dscr")),
            ("Debt/Equity", features.get("debt_equity_ratio")),
            ("Current Ratio", features.get("current_ratio")),
            ("GST-Banking Ratio", features.get("gst_banking_ratio")),
            ("Consistency Score", cross_validation.get("overall_data_consistency_score")),
        ]:
            row = table.add_row().cells
            row[0].text = str(metric)
            if metric == "Consistency Score":
                row[1].text = format_percentage(value)
            else:
                row[1].text = format_ratio(value)

    @staticmethod
    def _section_risk_assessment(
        doc: Document,  # type: ignore[reportGeneralTypeIssues]
        explanation: Dict[str, Any],
        cross_validation: Dict[str, Any],
    ) -> None:
        doc.add_heading("SECTION 5: RISK ASSESSMENT", level=1)
        doc.add_paragraph("Consolidated risk matrix based on cross-validation anomalies and model explanation:")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Risk"
        table.rows[0].cells[1].text = "Impact"
        table.rows[0].cells[2].text = "Mitigation"
        for risk in explanation.get("top_negative_factors", [])[:5]:
            row = table.add_row().cells
            row[0].text = risk
            row[1].text = "Medium/High"
            row[2].text = "Quarterly covenant tracking and early warning triggers."
        for anomaly in cross_validation.get("anomalies", [])[:6]:
            row = table.add_row().cells
            row[0].text = f"{anomaly.get('title', 'Anomaly')}: {anomaly.get('details', '')}"
            row[1].text = str(anomaly.get("severity", "MEDIUM"))
            row[2].text = "Targeted forensic validation and covenant tightening."
        for indicator in cross_validation.get("fraud_indicators", [])[:4]:
            row = table.add_row().cells
            row[0].text = f"{indicator.get('indicator', 'Fraud indicator')} (source: {indicator.get('source', 'N/A')})"
            row[1].text = str(indicator.get("severity", "HIGH"))
            row[2].text = "Escalate to enhanced monitoring and transaction-level verification."
        doc.add_paragraph(
            "Overall data consistency score: "
            f"{format_percentage(cross_validation.get('overall_data_consistency_score'))}"
        )

    @staticmethod
    def _section_research(
        doc: Document,  # type: ignore[reportGeneralTypeIssues]
        findings: List[Dict[str, Any]],
        research_narrative: str | None,
    ) -> None:
        doc.add_heading("SECTION 6: RESEARCH FINDINGS", level=1)
        critical = [f for f in findings if str(f.get("severity")) in {"CRITICAL", "HIGH"}]
        informational = [f for f in findings if str(f.get("severity")) not in {"CRITICAL", "HIGH"}]
        doc.add_paragraph(
            f"Total findings reviewed: {len(findings)} | "
            f"Critical/High: {len(critical)} | Other: {len(informational)}"
        )
        if research_narrative:
            doc.add_paragraph(research_narrative)
            doc.add_paragraph()
            doc.add_heading("Source Highlights", level=2)
        if critical:
            doc.add_heading("Critical/High Priority Alerts", level=2)
            for finding in critical[:6]:
                doc.add_paragraph(
                    f"[{finding.get('severity')}] {finding.get('source_name')}: "
                    f"{finding.get('summary')} (URL: {finding.get('source_url', 'N/A')})"
                )
        if informational:
            doc.add_heading("Contextual Findings", level=2)
            for finding in informational[:4]:
                doc.add_paragraph(
                    f"[{finding.get('severity')}] {finding.get('source_name')}: "
                    f"{finding.get('summary')} (URL: {finding.get('source_url', 'N/A')})"
                )

    @staticmethod
    def _section_recommendation(
        doc: Document,  # type: ignore[reportGeneralTypeIssues]
        decision: Dict[str, Any],
        explanation: Dict[str, Any],
        due_diligence: Dict[str, Any],
    ) -> None:
        doc.add_heading("SECTION 7: RECOMMENDATION", level=1)
        doc.add_paragraph(f"Final credit score: {decision.get('credit_score')}/900")
        doc.add_paragraph(f"Normalized score: {decision.get('normalized_score', 'N/A')}/100")
        doc.add_paragraph(f"Risk grade: {decision.get('risk_grade', 'N/A')}")
        doc.add_paragraph(f"Decision: {decision.get('recommendation')}")
        doc.add_paragraph(
            f"Recommended loan amount: {format_currency_cr(decision.get('recommended_loan_amount', 0))}"
        )
        doc.add_paragraph(
            f"Recommended interest rate: {decision.get('recommended_interest_rate', 0):.2f}% p.a."
        )
        doc.add_paragraph(
            f"Interest premium: +{int(decision.get('interest_premium_bps', 0) or 0)} bps over benchmark"
        )
        doc.add_paragraph("Tenor: 36 months")
        validity_days = 30
        decision_valid_till = date.fromordinal(date.today().toordinal() + validity_days)
        doc.add_paragraph(
            f"Decision validity window: {validity_days} days (valid up to {decision_valid_till.isoformat()})"
        )
        if due_diligence.get("due_diligence_risk_adjustment") not in (None, 0, 0.0):
            doc.add_paragraph(
                "Borrower interaction adjustment applied in scoring: "
                f"{float(due_diligence.get('due_diligence_risk_adjustment', 0.0)):+.1f}"
            )
        if decision.get("human_input_impact_points") not in (None, 0, 0.0):
            doc.add_paragraph(
                "Net score impact from human/borrower input: "
                f"{float(decision.get('human_input_impact_points', 0.0)):+.1f} points"
            )
        doc.add_paragraph("AI Explainability Note:")
        doc.add_paragraph(explanation.get("decision_narrative", "No narrative available"))
        if explanation.get("top_positive_factors"):
            doc.add_paragraph("Top positive factors:")
            for factor in explanation.get("top_positive_factors", [])[:3]:
                doc.add_paragraph(f"- {factor}")
        if explanation.get("top_negative_factors"):
            doc.add_paragraph("Top risk factors:")
            for factor in explanation.get("top_negative_factors", [])[:3]:
                doc.add_paragraph(f"- {factor}")

    @staticmethod
    def _section_mitigants(
        doc: Document,  # type: ignore[reportGeneralTypeIssues]
        cross_validation: Dict[str, Any],
        features: Dict[str, float],
    ) -> None:
        doc.add_heading("SECTION 8: RISK MITIGANTS & CONDITIONS", level=1)
        covenants: List[str] = [
            "Monthly stock and receivable statements to lender.",
            "Quarterly GST vs banking reconciliation with management certification.",
            "DSCR floor covenant at 1.20x; breach to trigger corrective action plan.",
        ]
        if float(features.get("gstr3b_vs_2a_itc_gap", 0.0) or 0.0) > 5:
            covenants.append("Monthly GST ITC reconciliation submission until mismatch normalizes below 5%.")
        if float(features.get("debt_equity_ratio", 0.0) or 0.0) > 1.5:
            covenants.append("No incremental unsecured debt without lender approval.")
        if cross_validation.get("fraud_indicators"):
            covenants.append("Enhanced transaction monitoring on related-party and round-amount transfers.")
        for covenant in covenants:
            doc.add_paragraph(f"- {covenant}")

    @staticmethod
    def _section_annexures(
        doc: Document,  # type: ignore[reportGeneralTypeIssues]
        findings: List[Dict[str, Any]],
        explanation: Dict[str, Any],
        due_diligence: Dict[str, Any],
        decision: Dict[str, Any],
    ) -> None:
        doc.add_heading("SECTION 9: ANNEXURES", level=1)
        doc.add_paragraph("Annexure A: Raw financial data tables")
        doc.add_paragraph("Annexure B: Research source URLs")
        for finding in findings[:20]:
            doc.add_paragraph(f"- {finding.get('source_url')}")
        doc.add_paragraph("Annexure C: SHAP explanation chart (placeholder)")
        doc.add_paragraph(f"Top factors: {', '.join(list(explanation.get('shap_waterfall_data', {}).keys())[:10])}")
        doc.add_paragraph("Annexure D: Glossary of Indian credit terms")
        doc.add_paragraph("Annexure E: Decision Traceability and Governance")
        doc.add_paragraph(f"CAM generation date: {date.today().isoformat()}")
        doc.add_paragraph("Model stack: Rules + ML calibration + explainability narrative")
        doc.add_paragraph(f"Total web-research findings considered: {len(findings)}")
        doc.add_paragraph(
            "Human-in-the-loop adjustment from due diligence: "
            f"{float(due_diligence.get('due_diligence_risk_adjustment', 0.0)):+.1f}"
        )
        doc.add_paragraph(
            "Final decision score delta from human input: "
            f"{float(decision.get('human_input_impact_points', 0.0)):+.1f} points"
        )

    @staticmethod
    def _add_footer_page_number(doc: Document) -> None:  # type: ignore[reportGeneralTypeIssues]
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = p.add_run("Page ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.addnext(fld)
