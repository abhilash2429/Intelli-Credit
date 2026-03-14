"""
Investment Report Generator.
Generates a structured downloadable DOCX covering all 4 hackathon stages:
entity profile, document analysis, secondary research, SWOT, reasoning, recommendation.
"""
from __future__ import annotations

import logging
import os
import uuid
from datetime import datetime

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

_BLUE = RGBColor(0x1E, 0x6F, 0xD9)
_INK = RGBColor(0x0F, 0x1F, 0x3D)
_GREEN = RGBColor(0x16, 0xA3, 0x4A)
_RED = RGBColor(0xDC, 0x26, 0x26)
_GRAY = RGBColor(0x6B, 0x72, 0x80)


async def generate_investment_report(company_id: str, db: AsyncSession) -> str:
    from backend.models.db_models import (
        Company,
        Document,
        DocumentClassification,
        LoanApplication,
        ResearchFindingRecord,
        RiskScore,
        SwotAnalysis,
    )

    company = await db.get(Company, uuid.UUID(company_id))
    if not company:
        raise ValueError(f"Company {company_id} not found")

    async def latest(model, filter_col):
        r = await db.execute(
            select(model)
            .where(filter_col == uuid.UUID(company_id))
            .order_by(model.created_at.desc())
            .limit(1)
        )
        return r.scalar_one_or_none()

    loan = await latest(LoanApplication, LoanApplication.company_id)
    score = await latest(RiskScore, RiskScore.company_id)
    swot = await latest(SwotAnalysis, SwotAnalysis.company_id)

    findings_r = await db.execute(
        select(ResearchFindingRecord).where(
            ResearchFindingRecord.company_id == uuid.UUID(company_id)
        )
    )
    findings = findings_r.scalars().all()

    clf_r = await db.execute(
        select(DocumentClassification)
        .where(DocumentClassification.company_id == uuid.UUID(company_id))
        .where(DocumentClassification.human_approved != False)  # noqa: E712
    )
    classifications = clf_r.scalars().all()

    # Build document
    doc = DocxDocument()
    _set_base_style(doc)

    # Cover
    _heading(doc, "INVESTMENT CREDIT REPORT", level=0)
    _info_table(
        doc,
        [
            ("Entity", company.name),
            ("CIN", getattr(company, "cin", None) or "N/A"),
            ("Sector", getattr(company, "sector", None) or "N/A"),
            ("PAN", getattr(company, "pan_number", None) or "N/A"),
            ("Report Date", datetime.now().strftime("%B %d, %Y")),
            ("Prepared by", "Intelli-Credit AI Engine v2.0"),
        ],
    )
    doc.add_page_break()

    # §1 Executive Summary
    _heading(doc, "§1 Executive Summary", 1)
    if score:
        decision_color = (
            _GREEN
            if score.decision == "APPROVE"
            else (_RED if score.decision == "REJECT" else None)
        )
        p = doc.add_paragraph()
        run = p.add_run(f"Credit Decision: {score.decision}")
        run.bold = True
        if decision_color:
            run.font.color.rgb = decision_color
        doc.add_paragraph(
            f"Risk Score: {score.final_risk_score:.1f} / 100 — {score.risk_category}"
        )
        if score.recommended_limit_crore:
            doc.add_paragraph(
                f"Recommended Limit: ₹{score.recommended_limit_crore:.1f} Cr"
            )
    if swot and swot.investment_thesis:
        doc.add_paragraph(f"Investment Thesis: {swot.investment_thesis}")
    doc.add_page_break()

    # §2 Entity & Loan Profile
    _heading(doc, "§2 Entity & Loan Profile", 1)
    entity_rows = [
        ("Company Name", company.name),
        ("CIN", getattr(company, "cin", None) or "N/A"),
        ("Sector", getattr(company, "sector", None) or "N/A"),
        (
            "Annual Turnover",
            f"₹{getattr(company, 'annual_turnover_cr', None) or 'N/A'} Cr",
        ),
        (
            "Year of Incorporation",
            str(getattr(company, "year_of_incorporation", None) or "N/A"),
        ),
    ]
    _info_table(doc, entity_rows)

    if loan:
        _heading(doc, "Loan Details", 2)
        _info_table(
            doc,
            [
                ("Loan Type", loan.loan_type),
                ("Amount Requested", f"₹{loan.loan_amount_cr:.1f} Cr"),
                ("Tenure", f"{loan.tenure_months} months"),
                ("Proposed Rate", f"{loan.proposed_rate_pct or 'TBD'}% p.a."),
                ("Repayment Mode", loan.repayment_mode or "N/A"),
                ("Purpose", loan.purpose or "N/A"),
                (
                    "Collateral",
                    f"{loan.collateral_type or 'None'} — ₹{loan.collateral_value_cr or 0:.1f} Cr",
                ),
            ],
        )
    doc.add_page_break()

    # §3 Documents Analysed
    _heading(doc, "§3 Documents Analysed & Classification", 1)
    for clf in classifications:
        etype = clf.human_type_override or clf.auto_type
        status_str = "✓ APPROVED" if clf.human_approved else "PENDING"  # type: ignore[reportGeneralTypeIssues]
        if clf.human_type_override:  # type: ignore[reportGeneralTypeIssues]
            status_str += f" (overridden from {clf.auto_type})"
        doc.add_paragraph(
            f"  • {etype} — {status_str} | Auto-confidence: {clf.auto_confidence:.0%} | {clf.auto_reasoning or ''}",
            style="List Bullet",
        )
    if not classifications:
        doc.add_paragraph("No classified documents available.")
    doc.add_page_break()

    # §4 Financial Analysis
    _heading(doc, "§4 Financial Analysis", 1)
    if score and score.rule_based_score is not None:
        _heading(doc, "Score Breakdown", 2)
        _info_table(
            doc,
            [
                ("Rule-based Score", f"{score.rule_based_score:.1f} / 100"),
                (
                    "ML Stress Probability",
                    f"{score.ml_stress_probability:.1%}"
                    if score.ml_stress_probability
                    else "N/A",
                ),
                ("Final Weighted Score", f"{score.final_risk_score:.1f} / 100"),
            ],
        )

    if score and score.shap_values:
        _heading(doc, "Top Risk Drivers (SHAP)", 2)
        shap_sorted = sorted(
            score.shap_values.items(), key=lambda x: abs(x[1]), reverse=True
        )[:8]
        for feat, val in shap_sorted:
            direction = "▲ RISK" if val > 0 else "▼ RISK"
            doc.add_paragraph(
                f"  • {feat}: {val:+.3f} ({direction})", style="List Bullet"
            )
    doc.add_page_break()

    # §5 Secondary Research
    _heading(doc, "§5 Secondary Research Findings", 1)
    if findings:
        for f in sorted(
            findings,
            key=lambda x: getattr(x, "severity", "LOW") if hasattr(x, "severity") else "LOW",
            reverse=True,
        ):
            severity = getattr(f, "severity", "LOW")
            summary = getattr(f, "summary", str(f))
            source = getattr(f, "source_name", "Web Research")
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{severity}] ")
            run.bold = True
            if severity == "HIGH":
                run.font.color.rgb = _RED
            elif severity == "MEDIUM":
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
            p.add_run(f"{summary} — Source: {source}")
    else:
        doc.add_paragraph("No adverse research findings identified.")

    if swot and swot.sector_outlook:
        _heading(doc, "Sector & Macro Context", 2)
        doc.add_paragraph(swot.sector_outlook)
        if swot.macro_signals:
            ms = swot.macro_signals
            _info_table(
                doc,
                [
                    ("RBI Repo Rate", f"{ms.get('rbi_repo_rate_pct', 'N/A')}%"),
                    ("India GDP Growth", f"{ms.get('india_gdp_growth_pct', 'N/A')}%"),
                    (
                        "Sector Credit Growth",
                        f"{ms.get('sector_credit_growth_pct', 'N/A')}%",
                    ),
                    ("CPI Inflation", f"{ms.get('inflation_cpi_pct', 'N/A')}%"),
                ],
            )
    doc.add_page_break()

    # §6 SWOT Analysis
    _heading(doc, "§6 SWOT Analysis", 1)
    if swot:
        for quadrant, label in [
            ("strengths", "STRENGTHS"),
            ("weaknesses", "WEAKNESSES"),
            ("opportunities", "OPPORTUNITIES"),
            ("threats", "THREATS"),
        ]:
            _heading(doc, label, 2)
            items = getattr(swot, quadrant) or []
            for item in items:
                if isinstance(item, dict):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(item.get("point", "")).bold = True
                    p.add_run(
                        f" — {item.get('evidence', '')} [{item.get('source', '')}]"
                    )
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph("SWOT analysis not yet generated.")
    doc.add_page_break()

    # §7 Reasoning Engine
    _heading(doc, "§7 Reasoning Engine — Why This Decision", 1)
    if score:
        violations = score.rule_violations or []
        if violations:
            _heading(doc, "Hard Stop Rules Triggered", 2)
            for v in violations:
                p = doc.add_paragraph(style="List Bullet")
                if isinstance(v, dict):
                    run = p.add_run(f"⚠ {v.get('id', '')}: ")
                    run.bold = True
                    run.font.color.rgb = _RED
                    p.add_run(v.get("description", str(v)))
                else:
                    run = p.add_run(f"⚠ {v}")
                    run.bold = True
                    run.font.color.rgb = _RED

        strengths = score.risk_strengths or []
        if strengths:
            _heading(doc, "Credit Strengths", 2)
            for st in strengths:
                if isinstance(st, dict):
                    doc.add_paragraph(
                        f"✓ {st.get('description', str(st))}", style="List Bullet"
                    )
                else:
                    doc.add_paragraph(f"✓ {st}", style="List Bullet")

        if score.decision_rationale:
            _heading(doc, "AI Narrative", 2)
            doc.add_paragraph(score.decision_rationale)
    doc.add_page_break()

    # §8 Recommendation
    _heading(doc, "§8 Recommendation", 1)
    if score:
        p = doc.add_paragraph()
        run = p.add_run(f"DECISION: {score.decision}")
        run.bold = True
        run.font.size = Pt(14)
        if loan:
            doc.add_paragraph(
                f"Requested: ₹{loan.loan_amount_cr:.1f} Cr | Recommended: ₹{score.recommended_limit_crore or loan.loan_amount_cr:.1f} Cr"
            )
        if getattr(score, "interest_premium_bps", None):
            doc.add_paragraph(
                f"Pricing Guidance: MCLR + {score.interest_premium_bps}bps"
            )
    if swot and swot.recommendation:
        doc.add_paragraph(swot.recommendation)

    # §9 Declaration
    _heading(doc, "§9 Declaration", 1)
    doc.add_paragraph(
        "This report has been generated by the Intelli-Credit AI Engine. "
        "All analysis is subject to review and approval by a qualified credit officer. "
        "The AI engine provides structured analysis and quantitative inputs; "
        "the final credit decision rests with the authorised credit authority per RBI guidelines."
    )
    doc.add_paragraph(
        f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} IST"
    )
    doc.add_paragraph(
        "Intelli-Credit | IIT Hyderabad × Vivriti Capital Hackathon 2026"
    )

    # Save
    output_dir = os.environ.get("CAM_OUTPUT_DIR", "outputs/reports")
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"investment_report_{company_id}.docx")
    doc.save(path)
    logger.info(f"Investment report saved: {path}")
    return path


def _heading(doc, text: str, level: int):
    if level == 0:
        p = doc.add_heading(text, 0)
    else:
        p = doc.add_heading(text, level)
    for run in p.runs:
        run.font.color.rgb = _INK


def _info_table(doc, rows: list):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        cell_label.text = str(label)
        cell_value.text = str(value)
        for run in cell_label.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = _GRAY
    doc.add_paragraph("")


def _set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
